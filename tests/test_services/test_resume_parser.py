import pytest
from backend.services.resume_parser import ResumeParser


@pytest.fixture
def parser():
    """Provides a resume parser instance"""
    return ResumeParser()


def test_extract_contact_email(parser):
    """Test email extraction"""
    text = "Contact me at john.doe@example.com for more information"
    contact = parser.extract_contact(text)
    
    assert contact["email"] == "john.doe@example.com"


def test_extract_contact_phone(parser):
    """Test phone number extraction"""
    text = "Call me at (555) 123-4567 or 555-987-6543"
    contact = parser.extract_contact(text)
    
    assert contact["phone"] is not None
    assert "555" in contact["phone"]


def test_extract_contact_both(parser):
    """Test extraction of both email and phone"""
    text = """
    John Doe
    Email: john@example.com
    Phone: 123-456-7890
    """
    contact = parser.extract_contact(text)
    
    assert contact["email"] == "john@example.com"
    assert "123" in contact["phone"]


def test_extract_years_experience_pattern1(parser):
    """Test years extraction - pattern: 'X years of experience'"""
    text = "Software engineer with 5 years of experience in Python"
    years = parser.extract_years_experience(text)
    
    assert years == 5


def test_extract_years_experience_pattern2(parser):
    """Test years extraction - pattern: 'X+ years experience'"""
    text = "Looking for candidates with 7+ years experience"
    years = parser.extract_years_experience(text)
    
    assert years == 7


def test_extract_years_experience_none(parser):
    """Test when no years mentioned"""
    text = "Fresh graduate looking for opportunities"
    years = parser.extract_years_experience(text)
    
    assert years is None


@pytest.mark.asyncio
async def test_parse_docx_basic(parser):
    """Test basic DOCX parsing structure"""
    # Create minimal DOCX in memory
    from docx import Document
    from io import BytesIO
    
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Python Developer")
    doc.add_paragraph("Email: test@example.com")
    doc.add_paragraph("Skills: Python, FastAPI, Docker")
    
    # Save to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    # Parse
    text = await parser.parse_docx(docx_bytes.read())
    
    assert "John Doe" in text
    assert "Python Developer" in text
    assert "test@example.com" in text


@pytest.mark.asyncio
async def test_parse_invalid_extension(parser):
    """Test error handling for invalid file type"""
    with pytest.raises(ValueError) as exc_info:
        await parser.parse(b"some bytes", "resume.txt")
    
    assert "Unsupported file format" in str(exc_info.value)